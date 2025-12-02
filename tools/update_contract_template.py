"""
Update the Contract of Sale template with docxtemplater placeholders.
Creates a new file contacts/Contract Template.docx preserving the original.
"""

from pathlib import Path
from docx import Document

SOURCE_PATH = Path("contacts/Contract of Sale.docx")
TARGET_PATH = Path("contacts/Contract Template.docx")


def replace_paragraph_contains(doc: Document, needle: str, replacement: str) -> None:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = replacement
            return
    raise ValueError(f"Could not find paragraph containing: {needle!r}")


def replace_text_in_paragraph(doc: Document, needle: str, replacement: str) -> None:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = paragraph.text.replace(needle, replacement)


def main() -> None:
    doc = Document(SOURCE_PATH)

    replace_paragraph_contains(
        doc,
        "This Agreement dated",
        (
            "This Agreement dated {agreementDate} is between (Buyer: {buyerName}, "
            "{buyerFullAddress}, {buyerPhone}, {buyerEmail}) herein referred to as Buyer "
            "and {breederName} of {kennelName} herein referred to as Breeder."
        ),
    )

    replace_paragraph_contains(
        doc,
        "In Consideration of the total sum",
        (
            "In Consideration of the total sum of {salePrice} ({salePriceWords}) and the mutual promises "
            "contained herein, Breeder has agreed to sell, and Buyer has agreed to purchase "
            "{puppyCount} ({maleCount} male, {femaleCount} female) American Bully puppy."
        ),
    )

    replace_paragraph_contains(doc, "Born on", "Born on {puppyDOBLong}")
    replace_paragraph_contains(doc, "Sire:", "Sire: {sireName}")
    replace_paragraph_contains(doc, "Dam:", "Dam: {damName}")

    replace_paragraph_contains(
        doc,
        "The puppy is sold as a pet",
        "{#isPet}The puppy is sold as a pet, with no registration, and must be spayed/neutered at no earlier than 18 months of age no later than two years of age (as early spay/neuter can be detrimental to the dogs overall health and wellness).{/isPet}",
    )

    replace_paragraph_contains(
        doc,
        "The puppy is sold with breeding rights",
        "{#isFullRights}The puppy is sold with breeding rights, or “Full Rights\" registration. Buyer will make a good faith effort to show the dog or allow the dog to be shown by the Breeder, to its ABKC and UKC Championship.{/isFullRights}",
    )

    replace_paragraph_contains(
        doc,
        "If No Registration",
        "{#isPet}If \"No Registration\" has been selected, this puppy is being sold as pet quality only, intended for companionship with no guarantees as to breeding soundness, show-ability, work ability, trainability, temperament or size at maturity. The Buyer agrees to NO BREEDING of this dog (accidental or intentional).  Buyer is to have the dog altered (Spay/Neuter/Vasectomy/Hysterectomy) by the age of 2 years old (24 months) but no earlier than 18 months of age.{/isPet}",
    )

    replace_paragraph_contains(
        doc,
        "The Buyer affirms that their purchase",
        "{#isPet}The Buyer affirms that their purchase is for a \"pet home\" only and not for breeding purposes. If the dog produces a litter without the knowledge and written consent of the Breeder, the Breeder will be entitled to compensation in the amount of $5,000 (Five Thousand Dollars and no Cents) for breach of contract terms. The Buyer herein agrees to pay the Breeder an additional $2,000 (Two Thousand Dollars and no cents) per puppy produced (dead or alive) from the whelping no later than 30 days after the birth of the unwarranted breeding.{/isPet}",
    )

    replace_paragraph_contains(
        doc,
        "The General Health Guarantee",
        "{#isPet}The General Health Guarantee also becomes null and void if spay/neuter contract is violated.{/isPet}",
    )

    replace_text_in_paragraph(doc, "State of ________, County of ________", "State of {state}, County of {county}")
    replace_text_in_paragraph(
        doc,
        "State of ___________ County of ____________",
        "State of {state} County of {county}",
    )
    replace_text_in_paragraph(
        doc,
        "State of (______), County of (_____)",
        "State of ({state}), County of ({county})",
    )

    replace_paragraph_contains(doc, "Signed on", "Signed on {signingDate}")

    replace_paragraph_contains(
        doc,
        "On this _____ day of _________",
        "On this {signingDate}, before me, the undersigned, a Notary Public in and for said State, personally appeared _____________________, personally known to me or proved to me on the basis of satisfactory evidence to be the individual whose name is subscribed to the within Instrument and acknowledged to me that s/he/they executed the same in her/his/their capacity, and that by her/his/their signature on the instrument, the individuals, or the person upon behalf of which the individuals acted, executed the instrument.",
    )

    replace_paragraph_contains(
        doc,
        "This Agreement is made and entered into this ______ day of",
        "This Agreement is made and entered into this {agreementDate}",
    )
    replace_paragraph_contains(
        doc,
        "By and between",
        "By and between {breederName} (Breeder) and {buyerName} (Buyer),",
    )
    replace_paragraph_contains(
        doc,
        "For the purpose of setting forth the terms",
        "For the purpose of setting forth the terms and conditions of purchase by the Buyer of a Purebred American Bully from the litter born on {puppyDOBLong}. Out of {sireName} (Sire), and {damName} (Dam). "
        "For {salePrice} the Breeder agrees to sell and buyer agrees to purchase a {femaleCount} female, {maleCount} male companion puppy from the litter described above subject to the following terms.",
    )

    for paragraph in doc.paragraphs:
        normalized = paragraph.text.replace("\xa0", " ")
        stripped = normalized.strip()
        if stripped.startswith("STATE OF") and ")" in stripped:
            suffix_index = stripped.find(")")
            suffix = stripped[suffix_index:]
            paragraph.text = f"STATE OF {{state}} {suffix}"
        if stripped.startswith("COUNTY") and ")SS" in stripped:
            paragraph.text = "COUNTY OF {county} )SS.:"

    doc.save(TARGET_PATH)
    print(f"Wrote updated template to {TARGET_PATH}")


if __name__ == "__main__":
    main()

